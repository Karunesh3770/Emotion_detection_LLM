import speech_recognition as sr
from docx import Document
from transformers import AutoTokenizer, AutoModelForSequenceClassification
import torch.nn.functional as F
import torch, re
from datetime import datetime
import threading

# Load model once
model_name = "bhadresh-savani/bert-base-go-emotion"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForSequenceClassification.from_pretrained(model_name)

# Prepare document
doc_path = r"C:\Users\Dell\OneDrive\Desktop\ergolab\ergo refined project\speech2.docx"
doc = Document()
doc.add_paragraph("Speech & Emotion Log\n")

# Flag to stop recording
stop_flag = False

# Listen for 'q' in the background
def listen_for_quit():
    global stop_flag
    input("Press ENTER after speaking your full paragraph (or type 'q' to stop): ")
    stop_flag = True

threading.Thread(target=listen_for_quit, daemon=True).start()

# Setup recognizer
r = sr.Recognizer()

print("\n🎙 Please speak a full paragraph now. Recording until you press ENTER...\n")

# Record full paragraph (non-stop)
try:
    with sr.Microphone() as source:
        r.adjust_for_ambient_noise(source)
        print("🎧 Listening now... speak your full paragraph.")
        audio = r.listen(source, timeout=5, phrase_time_limit=60)  # up to 60 seconds

    # Transcribe
    try:
        text = r.recognize_google(audio)
        print("\n You said:\n", text)
    except sr.UnknownValueError:
        print(" Could not understand audio.")
        text = "[Unrecognized speech]"
    except sr.RequestError as e:
        print(f"API request error: {e}")
        text = "[API Error]"

    # Clean text
    cleaned = re.sub(r'\s+', ' ', text).strip()

    # Emotion detection
    if cleaned:
        inputs = tokenizer(cleaned, return_tensors="pt", truncation=True, padding=True)
        outputs = model(**inputs)
        probs = F.softmax(outputs.logits, dim=1)
        pred_class = torch.argmax(probs, dim=1).item()
        labels = model.config.id2label
        emotion = labels[pred_class]
    else:
        emotion = "N/A"

    # Log to Word
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"\nTimestamp: {timestamp}")
    doc.add_paragraph(f"🗣 Full Paragraph: {text}")
    doc.add_paragraph(f" Detected Emotion: {emotion}")
    doc.save(doc_path)

    print(f"\nEmotion Detected: {emotion}")
    print(f"Logged in: {doc_path}")

except Exception as e:
    print(f" Error: {e}")

print(" Session ended.")
